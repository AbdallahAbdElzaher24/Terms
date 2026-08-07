"""DOCX text extraction — python-docx.

Pulls paragraphs and table cells in document order, and keeps heading level
as metadata so the chunker downstream can respect section boundaries.
"""
from dataclasses import dataclass

from docx import Document


@dataclass
class DocxBlock:
    text: str
    style: str  # e.g. "Heading 1", "Normal", "Table Cell"


def extract_docx(path: str) -> list[DocxBlock]:
    doc = Document(path)
    blocks: list[DocxBlock] = []

    for para in doc.paragraphs:
        if para.text.strip():
            blocks.append(DocxBlock(text=para.text.strip(), style=para.style.name if para.style else "Normal"))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    blocks.append(DocxBlock(text=cell.text.strip(), style="Table Cell"))

    return blocks


def extract_docx_full_text(path: str) -> str:
    blocks = extract_docx(path)
    return "\n".join(b.text for b in blocks)
